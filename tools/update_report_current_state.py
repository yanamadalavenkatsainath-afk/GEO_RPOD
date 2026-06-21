from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\Venkat\OneDrive\Desktop\appex\flight sim")
DOCX = ROOT / "GEO_RPOD_Report_Current_State.docx"
PLOTS = ROOT / "monte_carlo_plots.png"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], style: str = "Table Grid") -> Paragraph:
    table = paragraph._parent.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.8))
    try:
        table.style = style
    except KeyError:
        pass
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            if r == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    paragraph._p.addnext(table._tbl)
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    spacer = Paragraph(new_p, paragraph._parent)
    return spacer


def set_text(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = bold


def find_para(doc: Document, exact: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    raise ValueError(f"Could not find paragraph: {exact}")


def find_para_index(doc: Document, exact: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == exact:
            return i
    raise ValueError(f"Could not find paragraph: {exact}")


def add_bullets_after(anchor: Paragraph, bullets: list[str]) -> Paragraph:
    current = anchor
    for bullet in bullets:
        current = insert_paragraph_after(current, bullet)
        try:
            current.style = "List Bullet"
        except Exception:
            pass
    return current


def main() -> None:
    doc = Document(DOCX)

    doc.core_properties.title = "GEO RPOD Simulation Current State Report"
    doc.core_properties.author = "Venkata Sainath Yanamadala"
    doc.core_properties.subject = "Autonomous GEO RPOD simulation current baseline, Monte Carlo results, and full-6DOF readiness"

    set_text(doc.paragraphs[0], "GEO RPOD Simulation — Current Full-6DOF Readiness Report", True)
    set_text(
        doc.paragraphs[1],
        "Autonomous GNC Stack for GEO Rendezvous, Proximity Operations, Terminal Docking and Capture Simulation",
    )
    set_text(doc.paragraphs[2], "Venkata Sainath Yanamadala | MSc Space Engineering, University of Surrey | May 2026")

    abstract_i = find_para_index(doc, "Abstract")
    set_text(
        doc.paragraphs[abstract_i + 1],
        (
            "This report documents the current state of the GEO RPOD simulation stack after the latest terminal "
            "docking, soft/hard capture, finite-body, contact-dynamics and full-6DOF readiness updates. The current "
            "software is a closed-loop simulation prototype: it propagates GEO relative motion, estimates deputy "
            "and target-relative state, commands RPOD guidance, tracks a simulated docking port, applies terminal "
            "capture gates, and records Monte Carlo stress-case performance. It is not yet a raw-image, flight-qualified "
            "uncooperative docking system."
        ),
    )
    set_text(
        doc.paragraphs[abstract_i + 2],
        (
            "The May 2026 revision distinguishes validated baseline behaviour from newly added switchable realism "
            "modules. The validated baseline still uses synthetic camera/range-style measurements and a known target "
            "body/port model. Newly integrated modules include physical thruster allocation, finite body geometry, "
            "coupled contact response, body-mounted camera field-of-view checks, moving keep-out zones, spin-sync "
            "control scaffolding, angles-only/NMC navigation hooks, and an interface for future uncooperative pose "
            "estimation from imagery."
        ),
    )

    replacements = {
        "14. Monte Carlo Validation — Version 2 (99.7%)": "14. Monte Carlo Validation — Historical Version 2 Baseline",
        "16. Current Status and Remaining Work": "16. Current Project Status and Full-6DOF Readiness",
        "14.  Monte Carlo Validation — Version 2 (99.7%)": "14.  Monte Carlo Validation — Historical Version 2 Baseline",
        "16.  Current Project Status": "16.  Current Project Status and Full-6DOF Readiness",
    }
    for p in doc.paragraphs:
        if p.text.strip() in replacements:
            set_text(p, replacements[p.text.strip()], True)

    # Add a current baseline note after the historical Monte Carlo heading.
    hist = find_para(doc, "14.  Monte Carlo Validation — Historical Version 2 Baseline")
    current = insert_paragraph_after(
        hist,
        (
            "Note: Sections 14.1-14.4 preserve the earlier validation narrative as historical context. The active "
            "May 2026 software baseline has since added terminal capture physics and full-6DOF readiness hooks; its "
            "latest 20-run Monte Carlo summary is reported in Section 16."
        ),
    )
    current.runs[0].italic = True

    completed = find_para(doc, "16.1  Completed Work")
    p = insert_paragraph_after(
        completed,
        (
            "The current baseline is a working closed-loop GEO RPOD simulator. It models GEO relative motion, deputy "
            "attitude control, target-relative navigation, terminal port tracking, approach-cone gating, soft capture "
            "and hard-capture confirmation. The simulation can repeatedly close from hundreds of metres to centimetre-level "
            "port error with low final relative velocity, and it records detailed Monte Carlo metrics for docking time, "
            "delta-V, port range, port velocity and stress-case pass rate."
        ),
    )
    p = insert_paragraph_after(
        p,
        (
            "The latest single-run example docked successfully at t = 7820.2 s (2.17 hr). Soft capture occurred at "
            "0.9 cm port range and 1.2 mm/s relative velocity, followed by hard capture at 1.4 cm and 2.1 mm/s. "
            "The contact abstraction reported J = 0.063 N s and severity = 0.02, with total delta-V = 4.7195 m/s "
            "and propellant = 109.22 g."
        ),
    )

    p = insert_paragraph_after(p, "16.1.1  Current Data Flow", None)
    p.runs[0].bold = True
    p = add_bullets_after(
        p,
        [
            "Orbit and relative dynamics feed the truth LVLH deputy-chief state into the navigation and guidance layers.",
            "Deputy attitude sensors feed the MEKF; the MEKF output gates entry from detumble and sun acquisition into fine pointing.",
            "Synthetic camera/range-style target measurements feed the TH-EKF and port tracker; these replace raw imagery in the current baseline.",
            "The TH-EKF state, chief attitude estimate and port geometry feed the RPOD controller and terminal guidance law.",
            "The terminal controller commands desired acceleration; the plant integrates deputy motion and accumulates delta-V.",
            "Capture logic evaluates port range, relative speed, cone/lateral margins, attitude alignment and contact severity before hard capture.",
        ],
    )

    p = insert_paragraph_after(p, "16.1.2  What Is Working Well", None)
    p.runs[0].bold = True
    p = add_bullets_after(
        p,
        [
            "The rendezvous and proximity phase is stable: the controller reliably reduces range from roughly 385-425 m at RPOD handover to the terminal corridor.",
            "The terminal phase now separates COM range from port range, so the simulation can reason about docking geometry instead of only chief-centre distance.",
            "Soft capture and hard capture are no longer a single Boolean event; the simulation holds, damps and confirms capture before declaring docking.",
            "The latest 20-run Monte Carlo set docked 20/20 cases across camera dropout, gyro bias, high pose noise, nominal, range dropout and weak-thruster stress cases.",
            "The system now exposes the important remaining realism switches, allowing controlled A/B testing before enabling full-6DOF effects in Monte Carlo.",
        ],
    )

    remaining = find_para(doc, "16.2  Remaining Work — Prioritised")
    p = insert_paragraph_after(
        remaining,
        (
            "The next work is not basic rendezvous. The remaining risk is the final metres of uncooperative docking: "
            "real visual pose estimation, observability without range, collision-free replanning around moving appendages, "
            "spin synchronisation, and validated post-contact dynamics. These items must be enabled and tested one at a time."
        ),
    )
    p = add_bullets_after(
        p,
        [
            "Enable and tune physical thruster allocation, then run the main scenario and a small Monte Carlo comparison.",
            "Enable body-mounted camera FOV and finite body geometry, then verify camera loss/recovery and clearance behaviour.",
            "Enable moving keep-out avoidance and inspect whether approach times or terminal delta-V increase under appendage stress.",
            "Enable coupled contact dynamics and confirm that impact impulse and post-capture angular-rate changes remain bounded.",
            "Enable spin-sync only after the above are stable, because it couples attitude control, sensing and translation.",
            "Replace synthetic pose measurements with a trained or dataset-backed uncooperative pose estimator before claiming real target autonomy.",
        ],
    )

    audit = find_para(doc, "16.3  Truth Dependency Audit — Current State")
    p = insert_paragraph_after(
        audit,
        (
            "The current baseline is still simulation-truth assisted. The target geometry, docking-port definition and measurement "
            "statistics are known to the simulator. The camera model produces measurement-like information rather than raw pixels. "
            "This is appropriate for GNC architecture development, but it is not yet a demonstration of uncooperative vision autonomy."
        ),
    )

    rows = [
        ["Module / Assumption", "Current State", "Flight-Real Gap"],
        ["Target geometry and docking port", "Known body model and defined port", "Unknown or partially known target geometry must be estimated from imagery/lidar."],
        ["Vision measurement source", "Synthetic camera/range-style updates", "Raw monocular/stereo/lidar pose pipeline trained and validated on space imagery."],
        ["Relative navigation", "TH-EKF with range/camera-style measurements plus angles-only hook", "Angles-only observability and NMC manoeuvres must be validated without range."],
        ["Terminal control", "Port-relative guidance, cone/lateral gates, soft/hard capture", "Must be tested with body-mounted FOV, actuator layout, keep-outs and contact enabled."],
        ["Contact dynamics", "Soft/hard capture abstraction plus coupled contact module", "Needs high-fidelity rigid-body/contact validation and post-capture stabilization proof."],
        ["Full-6DOF realism flags", "Integrated but default-off for baseline comparison", "Enable one at a time and re-run Monte Carlo before claiming final performance."],
    ]
    p = insert_table_after(p, rows)

    p = insert_paragraph_after(p, "16.4  Latest 20-Run Monte Carlo Baseline — May 2026", None)
    p.runs[0].bold = True
    p = insert_paragraph_after(
        p,
        (
            "The current folder summary reports n = 20 trials and n = 20 docked. This is a useful smoke-test baseline, "
            "not the final 300-run statistical validation. The main result is that the new capture stack does not break "
            "basic docking, but terminal delta-V still has a large tail and must be tuned before the next investor-grade result."
        ),
    )
    mc_rows = [
        ["Metric", "Latest 20-run value", "Interpretation"],
        ["Docking success", "20 / 20 (100.0%)", "All current smoke-test cases reached docking."],
        ["Mean / median total delta-V", "7.239 m/s / 4.948 m/s", "Median is reasonable; mean is pulled up by high terminal-hover cases."],
        ["95th percentile total delta-V", "16.743 m/s", "The long tail is the main current tuning target."],
        ["Mean / median time to dock", "2.632 hr / 2.174 hr", "Most runs dock in a few hours; outliers remain."],
        ["Final port range", "mean 2.6 cm, median 1.5 cm, 95th 5.4 cm", "Capture geometry is tight in successful cases."],
        ["Final port relative velocity", "mean 4 mm/s, median 2 mm/s, 95th 19 mm/s", "Soft capture is generally slow, with some faster terminal cases."],
        ["Mean / 95th propellant", "167.1 g / 386.3 g", "Acceptable for smoke test, but high-DV tail needs reduction."],
    ]
    p = insert_table_after(p, mc_rows)

    if PLOTS.exists():
        p = insert_paragraph_after(p, "Figure 16-1. Latest 20-run Monte Carlo diagnostic plot set.", None)
        p.runs[0].italic = True
        pic_p = insert_paragraph_after(p, "")
        pic_p.add_run().add_picture(str(PLOTS), width=Inches(6.8))
        p = pic_p

    p = insert_paragraph_after(p, "16.5  Full-6DOF Readiness Modules Added", None)
    p.runs[0].bold = True
    flag_rows = [
        ["Capability", "Implementation", "Baseline Flag"],
        ["Physical thruster layout", "plant/thruster_layout.py allocator for a bounded 16-thruster box layout", "ENABLE_PHYSICAL_THRUSTER_LAYOUT = False"],
        ["Finite body geometry", "plant/finite_body.py box-body clearance/collision checks", "ENABLE_FINITE_BODY_COLLISION = False"],
        ["Coupled contact dynamics", "plant/contact_dynamics.py impulse and angular-rate response", "ENABLE_COUPLED_CONTACT_DYNAMICS = False"],
        ["Body-mounted camera FOV", "sensors/body_camera.py visibility cone and line-of-sight checks", "ENABLE_BODY_MOUNTED_CAMERA_FOV = False"],
        ["Moving keep-out zones", "control/keepout_planner.py rotating appendage keep-out scaffold", "ENABLE_KEEP_OUT_AVOIDANCE = False"],
        ["Spin synchronisation", "control/spin_sync_controller.py body-rate matching scaffold", "ENABLE_SPIN_SYNC = False"],
        ["Angles-only/NMC navigation", "estimation/th_ekf.py bearing update and control/nmc_guidance.py arc generator", "Available hook, not active in baseline"],
        ["Uncooperative pose interface", "sensors/uncooperative_pose_sensor.py surrogate measurement interface", "No dataset-backed neural estimator yet"],
    ]
    p = insert_table_after(p, flag_rows)

    p = insert_paragraph_after(p, "16.6  Investor-Ready Framing", None)
    p.runs[0].bold = True
    p = insert_paragraph_after(
        p,
        (
            "The correct external statement is: the project has a working autonomous GEO RPOD simulation prototype that "
            "solves the closed-loop guidance/control problem in a high-fidelity software environment and now contains the "
            "interfaces needed for flight-real perception, actuator, keep-out and contact realism. The remaining work is "
            "to replace synthetic perception with dataset-backed pose estimation, enable the full-6DOF modules in controlled "
            "Monte Carlo campaigns, and validate contact/post-capture stability under harder uncooperative target cases."
        ),
    )

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
